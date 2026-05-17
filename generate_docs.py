#!/usr/bin/env python3
"""Generate comprehensive Word document for WF Enterprise Analytics Platform"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, fill_color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create document
doc = Document()

# Title
title = doc.add_heading('Wells Fargo Enterprise Credit Analytics Platform', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(215, 30, 40)

subtitle = doc.add_paragraph('Complete Application Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.italic = True

doc.add_paragraph('Version 1.0.0 | May 2026')
doc.add_paragraph()

# TABLE OF CONTENTS
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. Application Overview',
    '2. Key Features',
    '3. Technology Stack',
    '4. Project Structure',
    '5. Architecture & Design Patterns',
    '6. API Endpoints Documentation',
    '7. Database & Data Layer',
    '8. Frontend Components',
    '9. Code Explanations',
    '10. Setup & Deployment',
    '11. Configuration',
    '12. Troubleshooting',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 1. APPLICATION OVERVIEW
doc.add_heading('1. Application Overview', 1)
doc.add_heading('Purpose', 2)
doc.add_paragraph(
    'The Wells Fargo Enterprise Credit Analytics Platform is a comprehensive web application for managing and '
    'analyzing credit facility portfolios. It provides multi-level drill-down capabilities from facilities down to '
    'analyst comments, enabling credit analysts to gain deep insights into obligor relationships, transactions, and risks.'
)

doc.add_heading('Core Functionality', 2)
features = [
    'View and manage credit facility portfolio across the enterprise',
    'Drill down from Facilities → Obligors → Transactions → Comments',
    'Search, filter, and paginate across all data levels',
    'Export data in CSV, Excel, and Parquet formats',
    'Real-time KPI strips showing Active/Total counts',
    'Responsive grid with column customization',
    'Modal-based drill-down hierarchy with breadcrumb trails',
    'Wells Fargo branded UI with professional styling',
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# 2. KEY FEATURES
doc.add_heading('2. Key Features', 1)
doc.add_heading('Multi-Level Drill-Down', 2)
doc.add_paragraph('The application supports a 4-level hierarchical navigation:')
for item in ['Level 1: Credit Facilities', 'Level 2: Obligors', 'Level 3: Transactions', 'Level 4: Comments']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Data Export', 2)
doc.add_paragraph('Export functionality at every level in three formats:')
for fmt in ['CSV (Comma-Separated Values)', 'Excel (.xlsx)', 'Parquet (Apache columnar format)']:
    doc.add_paragraph(fmt, style='List Bullet')

doc.add_heading('Standalone Pages', 2)
doc.add_paragraph('In addition to drill-down workflows, standalone pages are provided for:')
for page in ['Obligors Page - View all obligors with transaction drill-down',
             'Transactions Page - View all transactions with comment drill-down']:
    doc.add_paragraph(page, style='List Bullet')

doc.add_page_break()

# 3. TECHNOLOGY STACK
doc.add_heading('3. Technology Stack', 1)

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Technology'
for cell in hdr:
    shade_cell(cell, 'D71E28')
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

stack = [
    ('Backend', 'Python 3.10+ with Flask 2.3'),
    ('Frontend', 'HTML5, CSS3, Vanilla JavaScript (ES6+)'),
    ('Data Grid', 'AG Grid Community 31.3.2 (local vendor files)'),
    ('Data Storage', 'CSV files with Pandas DataFrame caching'),
    ('Export', 'Pandas + openpyxl (Excel) + pyarrow (Parquet)'),
    ('Web Framework', 'Flask with Blueprint architecture'),
    ('HTTP Client', 'Fetch API with ApiUtils wrapper'),
    ('Styling', 'Custom CSS with Wells Fargo color palette'),
]

for layer, tech in stack:
    row = table.add_row().cells
    row[0].text = layer
    row[1].text = tech

doc.add_page_break()

# 4. PROJECT STRUCTURE
doc.add_heading('4. Project Structure', 1)
structure = '''SampleApp4/
├── run.py                    Flask entry point
├── requirements.txt          Dependencies
├── app/
│   ├── __init__.py           Flask factory
│   ├── config.py             Configuration
│   ├── controllers/          Route handlers
│   │   ├── main_controller.py
│   │   ├── api_controller.py
│   │   └── export_controller.py
│   ├── services/             Business logic
│   │   ├── facility_service.py
│   │   ├── obligor_service.py
│   │   ├── transaction_service.py
│   │   └── export_service.py
│   ├── repositories/         Data access
│   ├── models/               Domain models
│   ├── templates/            HTML templates
│   │   ├── base.html
│   │   ├── dashboard.html
│   │   ├── obligors.html
│   │   ├── transactions.html
│   │   └── components/
│   ├── static/               Client assets
│   │   ├── css/              Stylesheets
│   │   ├── js/               JavaScript modules
│   │   └── vendor/ag-grid/   Local AG Grid
│   └── data/                 CSV data files
└── .claude/settings.json     Claude Config'''
doc.add_paragraph(structure, style='Intense Quote')

doc.add_page_break()

# 5. ARCHITECTURE & DESIGN PATTERNS
doc.add_heading('5. Architecture & Design Patterns', 1)

doc.add_heading('Server-Side Architecture', 2)
patterns = [
    ('Blueprint Pattern', 'Flask Blueprints organize routes into modules'),
    ('Repository Pattern', 'Data access abstracted through repositories'),
    ('Service Layer', 'Business logic separated from data access'),
    ('Factory Pattern', 'Flask app created via factory function'),
    ('Dataclass Models', 'Domain objects with to_dict()/from_dict()'),
    ('Response Envelope', 'All responses: {"success": bool, "data": [...], "meta": {...}}'),
]
for pattern, desc in patterns:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(pattern + ': ').bold = True
    p.add_run(desc)

doc.add_heading('Client-Side Architecture', 2)
doc.add_paragraph('Frontend uses IIFE module pattern:')
for item in [
    'Each app (WFApp, ObligorsApp, TransactionsApp) is a self-contained IIFE',
    'Shared utilities (ApiUtils, GridManager, ModalManager, DrillDown) are global IIFEs',
    'No build step or ES modules; all JS served as-is',
    'AG Grid loaded from local vendor files',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Modal Drill-Down System', 2)
for item in [
    'Stack-based modal management with z-index calculated by depth',
    'Breadcrumb trails show full hierarchy',
    'Escape key closes only topmost modal',
    'CSS animations (300ms) with column sizing deferred to 320ms',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 6. API ENDPOINTS
doc.add_heading('6. API Endpoints Documentation', 1)

doc.add_heading('Response Format', 2)
doc.add_paragraph('All endpoints return JSON with consistent structure:')
doc.add_paragraph('{"success": true, "data": [...], "meta": {"total": 100, "page": 1, "per_page": 50}}',
                 style='Intense Quote')

doc.add_heading('Facility Endpoints', 2)
for ep in ['GET /api/facilities - Paginated facilities',
           'GET /api/facilities/<id>/obligors - Obligors for facility']:
    doc.add_paragraph(ep, style='List Bullet')

doc.add_heading('Obligor Endpoints', 2)
for ep in ['GET /api/obligors - All obligors',
           'GET /api/obligors/<id>/transactions - Transactions for obligor']:
    doc.add_paragraph(ep, style='List Bullet')

doc.add_heading('Transaction Endpoints', 2)
for ep in ['GET /api/transactions - All transactions',
           'GET /api/transactions/<id>/comments - Comments for transaction']:
    doc.add_paragraph(ep, style='List Bullet')

doc.add_heading('Export Endpoints', 2)
doc.add_paragraph('All support format=csv|excel|parquet:')
for exp in ['GET /api/export/facilities', 'GET /api/export/obligors', 'GET /api/export/transactions']:
    doc.add_paragraph(exp, style='List Bullet')

doc.add_page_break()

# 7. DATABASE & DATA LAYER
doc.add_heading('7. Database & Data Layer', 1)

doc.add_heading('Data Storage', 2)
doc.add_paragraph('CSV files loaded into pandas DataFrames on startup for fast, in-memory queries.')

doc.add_heading('CSV Files', 2)
for f in ['facilities.csv', 'obligors.csv', 'transactions.csv', 'comments.csv']:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('Data Enrichment', 2)
for item in ['Facilities: obligor_count', 'Obligors: transaction_count', 'Transactions: comment_count']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 8. FRONTEND COMPONENTS
doc.add_heading('8. Frontend Components & Pages', 1)

doc.add_heading('Dashboard Page', 2)
doc.add_paragraph('Main portfolio view with credit facilities grid, KPI strip, and obligor drill-down.')

doc.add_heading('Obligors Page', 2)
doc.add_paragraph('Standalone page listing all obligors with transaction drill-down.')

doc.add_heading('Transactions Page', 2)
doc.add_paragraph('Standalone page listing all transactions with comment drill-down.')

doc.add_page_break()

# 9. CODE EXPLANATIONS
doc.add_heading('9. Code Explanations', 1)

doc.add_heading('Flask Application Factory', 2)
doc.add_paragraph(
    'create_app() instantiates Flask, loads config, and registers blueprints. '
    'WERKZEUG_RUN_MAIN check prevents duplicate startup logs in debug mode.'
)

doc.add_heading('Service Layer', 2)
doc.add_paragraph(
    'Services orchestrate repositories and enrich data with counts. '
    'Example: get_all_obligors() searches, adds transaction_count, paginates, returns JSON-ready data.'
)

doc.add_heading('API Routes', 2)
doc.add_paragraph(
    'Routes extract pagination/search params, call services, wrap responses in error handlers. '
    'All return JSON envelope with success/data/metadata.'
)

doc.add_heading('GridManager (AG Grid Wrapper)', 2)
doc.add_paragraph(
    'Key fix: _onGridReady uses requestAnimationFrame() to defer sizeColumnsToFit() '
    'until browser layout is complete, preventing column clustering.'
)

doc.add_heading('ModalManager (Stack-Based)', 2)
doc.add_paragraph(
    'Maintains modal stack with z-index = 1100 + (depth-1)*100. '
    'Breadcrumb shows hierarchy with › separators.'
)

doc.add_heading('DrillDown Orchestrator', 2)
doc.add_paragraph(
    'Orchestrates modal creation, grid init, and data loading. '
    'Defers column sizing 320ms after mount to account for CSS animation.'
)

doc.add_page_break()

# 10. SETUP & DEPLOYMENT
doc.add_heading('10. Setup & Deployment Instructions', 1)

doc.add_heading('Local Development', 2)
doc.add_paragraph('1. Create virtual environment:')
doc.add_paragraph('python3 -m venv venv && source venv/bin/activate', style='Intense Quote')

doc.add_paragraph('2. Install dependencies:')
doc.add_paragraph('pip install -r requirements.txt', style='Intense Quote')

doc.add_paragraph('3. Generate sample data:')
doc.add_paragraph('python3 -c "from app.utils.data_generator import generate_all; generate_all(\'app/data\')"',
                 style='Intense Quote')

doc.add_paragraph('4. Start server:')
doc.add_paragraph('python3 run.py', style='Intense Quote')

doc.add_paragraph('5. Open in browser:')
doc.add_paragraph('http://localhost:5000', style='Intense Quote')

doc.add_heading('Production Deployment', 2)
doc.add_paragraph('Use Gunicorn WSGI server with Nginx reverse proxy:')
doc.add_paragraph('gunicorn -w 4 -b 0.0.0.0:8000 "app:create_app()"', style='Intense Quote')

doc.add_page_break()

# 11. CONFIGURATION
doc.add_heading('11. Configuration', 1)

doc.add_heading('Environment Variables', 2)
for e in ['FLASK_ENV - development|production', 'HOST - Server address (0.0.0.0)',
          'PORT - Server port (5000)', 'DATA_DIR - CSV path (app/data)']:
    doc.add_paragraph(e, style='List Bullet')

doc.add_page_break()

# 12. TROUBLESHOOTING
doc.add_heading('12. Troubleshooting', 1)

issues = [
    ('ERR_ADDRESS_INVALID', 'Use http://localhost:5000 instead of http://0.0.0.0:5000'),
    ('Columns clustered in grids', 'Ensure sizeColumnsToFit() + 320ms setTimeout in modal mount'),
    ('Modals behind each other', 'Z-index must be set via JavaScript, not CSS'),
    ('Sidebar links disabled', 'Check routes are correct (/obligors, /transactions)'),
    ('Duplicate startup logs', 'Guard with WERKZEUG_RUN_MAIN environment check'),
    ('Data not loading', 'Verify DATA_DIR path and CSV files exist'),
    ('Export fails', 'pip install openpyxl pyarrow'),
]

for issue, solution in issues:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(issue + ': ').bold = True
    p.add_run(solution)

# Save document
output_path = 'f:/Office/SampleApp4/WF_Enterprise_Analytics_Documentation.docx'
doc.save(output_path)
print("[SUCCESS] Word document created successfully!")
print(f"[INFO] Saved to: {output_path}")
